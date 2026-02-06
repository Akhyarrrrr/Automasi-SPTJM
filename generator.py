from __future__ import annotations

import io
import re
import zipfile
import tempfile
import subprocess
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
from slugify import slugify

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from excel_utils import detect_max_n, normalize_email


# =========================================================
# DATA
# =========================================================
@dataclass
class Person:
    nama: str
    nip: str
    fakultas: str
    rekening: str
    bank: str
    email: Optional[str]


@dataclass
class ResultRow:
    nama: str
    nip: str
    email: str
    status: str
    message: str


# =========================================================
# DOCX HELPERS
# =========================================================
def _style(p, *, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, justify=False):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else align
    for r in p.runs:
        r.font.name = "Cambria"
        r.font.size = Pt(size)
        r.bold = bold


def _format_tanggal(dt: datetime) -> str:
    return dt.strftime("%d %B %Y")


def _format_rupiah(val) -> str:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return ""
    if isinstance(val, (int, float)) and not pd.isna(val):
        x = int(val)
        return f"{x:,}".replace(",", ".")
    return str(val).strip()


def _safe_filename(nama: str, nip: str, ext: str) -> str:
    base = f"SPTJM_{slugify(nama)}_{nip}".strip("_")
    base = base[:120]
    return f"{base}.{ext.lstrip('.')}"


# =========================================================
# BORDER ONLY (TIDAK NGUBAH LAYOUT)
# =========================================================
def _hide_table_borders_keep_layout(table, mode: str = "nil"):
    """
    Hilangkan border tanpa mengubah layout tabel.
    mode:
      - "nil": benar-benar hilang
      - "white": border masih ada tapi warna putih
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")

        if mode == "white":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "FFFFFF")
        else:
            # nil = hilang total (paling bersih)
            el.set(qn("w:val"), "nil")

        tblBorders.append(el)

    tblPr.append(tblBorders)


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


# =========================================================
# LIBREOFFICE CONVERT
# =========================================================
def _find_soffice(soffice_path: str | None = None) -> str:
    if soffice_path:
        p = Path(soffice_path)
        if p.exists():
            return str(p)
        raise FileNotFoundError(
            f"Path soffice override tidak ditemukan: {soffice_path}"
        )

    for name in ["soffice.exe", "soffice.com", "soffice"]:
        p = shutil.which(name)
        if p and Path(p).exists():
            return p

    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
    ]
    for c in candidates:
        if Path(c).exists():
            return c

    raise FileNotFoundError(
        "LibreOffice tidak ditemukan. Isi manual path soffice di UI/Sidebar."
    )


def convert_docx_bytes_to_pdf_bytes(
    docx_bytes: bytes,
    *,
    soffice_path: str | None = None,
    timeout_sec: int = 120,
) -> bytes:
    soffice = _find_soffice(soffice_path)

    # lebih stabil pakai .exe
    if soffice.lower().endswith(".com"):
        alt = soffice[:-4] + ".exe"
        if Path(alt).exists():
            soffice = alt

    with tempfile.TemporaryDirectory(prefix="sptjm_lo_") as tmp:
        tmpdir = Path(tmp)

        lo_profile = tmpdir / "lo_profile"
        lo_profile.mkdir(parents=True, exist_ok=True)
        lo_profile_uri = lo_profile.as_uri()

        docx_path = tmpdir / "input.docx"
        docx_path.write_bytes(docx_bytes)

        cmd = [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            f"-env:UserInstallation={lo_profile_uri}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(tmpdir),
            str(docx_path),
        ]

        try:
            res = subprocess.run(
                cmd, capture_output=True, text=True, timeout=timeout_sec
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                f"LibreOffice timeout > {timeout_sec}s (convert DOCX->PDF)."
            )

        if res.returncode != 0:
            raise RuntimeError(
                "LibreOffice gagal convert DOCX->PDF.\n"
                f"cmd: {' '.join(cmd)}\n"
                f"stdout:\n{res.stdout}\n"
                f"stderr:\n{res.stderr}\n"
            )

        pdfs = list(tmpdir.glob("*.pdf"))
        if not pdfs:
            raise RuntimeError(
                "Konversi selesai tapi PDF tidak ditemukan.\n"
                f"stdout:\n{res.stdout}\n"
                f"stderr:\n{res.stderr}\n"
            )

        return pdfs[0].read_bytes()


# =========================================================
# BUILD DOCX (SAMA SEPERTI AWAL, CUMA BORDER STATEMENT HILANG)
# =========================================================
def _build_docx_from_scratch(
    person: Person,
    lampiran_rows: List[Dict[str, str]],
    exec_dt: datetime,
) -> Document:
    doc = Document()

    # ===== HALAMAN 1 =====
    p = doc.add_paragraph("SURAT PERNYATAAN TANGGUNGJAWAB MUTLAK (SPTJM)")
    _style(p, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph(
        "Biaya Subimt Artikel/Insentif Publikasi/Opini Media Massa/Hak Kekayaan Intelektual"
    )
    _style(p, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("")

    p = doc.add_paragraph("Yang bertanda tangan di bawah ini:")
    _style(p)

    # Identitas (tetap seperti awal)
    t = doc.add_table(rows=5, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(4.2)
    t.columns[1].width = Cm(9.0)

    id_rows = [
        ("Nama", f": {person.nama}"),
        ("NIP", f": {person.nip}"),
        ("Fakultas", f": {person.fakultas}"),
        ("Nomor Rekening", f": {person.rekening}"),
        ("Nama Bank", f": {person.bank}"),
    ]
    for r, (k, v) in zip(t.rows, id_rows):
        r.cells[0].text = k
        r.cells[1].text = v
        _style(r.cells[0].paragraphs[0])
        _style(r.cells[1].paragraphs[0])

    doc.add_paragraph("")

    p = doc.add_paragraph("Menyatakan dengan sesungguhnya bahwa:")
    _style(p)

    items = [
        "Biaya Submit Artikel yang saya ajukan seperti yang tersebut pada lampiran belum pernah saya pertanggungjawabkan pada penelitian yang telah dilaksanakan, atau belum pernah menerima bantuan publikasi dari pihak/sumber dana lainnya, dan jika di kemudian hari terbukti bahwa biaya submit artikel yang saya ajukan telah pernah menerima bantuan publikasi dari pihak/sumber dana lainnya, maka saya akan mengembalikan dana insentif yang saya terima ke rekening Universitas Syiah Kuala.",
        "Biaya submit artikel yang saya ajukan seperti yang tersebut pada lampiran belum pernah dipertanggungjawabkan pada laporan penelitian dan belum pernah menerima bantuan publikasi dari sumber dana lain. Apabila di kemudian hari terbukti sebaliknya, saya bersedia mengembalikan dana yang telah diterima ke rekening Universitas Syiah Kuala.",
        "Artikel ilmiah/opini media massa/hak kekayaan intelektual yang diajukan seperti yang tersebut pada lampiran bebas plagiarisme dan merupakan karya asli.",
        "Artikel ilmiah/opini media massa/hak kekayaan intelektual yang diajukan seperti yang tersebut pada lampiran belum pernah menerima insentif pada periode sebelumnya maupun dari sumber dana lain.",
        "Saya bersedia mengembalikan dana insentif apabila di kemudian hari terbukti bahwa karya yang diajukan bukan milik saya, sudah pernah menerima insentif, atau tidak sesuai dengan ketentuan yang berlaku.",
        "Nomor rekening dan nama bank yang saya cantumkan benar dan aktif untuk menerima dana insentif."
    ]

    # ✅ STATEMENT TABLE: persis seperti awal (kolom kiri kecil, kanan lebar)
    stmt_tbl = doc.add_table(rows=len(items), cols=2)
    stmt_tbl.autofit = False
    stmt_tbl.style = "Table Grid"
    stmt_tbl.columns[0].width = Cm(0.5)
    stmt_tbl.columns[1].width = Cm(14.7)

    for i, txt in enumerate(items, start=1):
        c_no = stmt_tbl.rows[i - 1].cells[0]
        c_tx = stmt_tbl.rows[i - 1].cells[1]
        c_no.text = str(i)
        c_tx.text = txt
        _style(c_no.paragraphs[0], size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        _style(c_tx.paragraphs[0], size=11, justify=True)

    # ✅ CUMA HILANGKAN BORDER, TANPA UBAH LAYOUT
    _hide_table_borders_keep_layout(stmt_tbl, mode="nil")
    # kalau kamu lebih suka "putih" daripada nil, ganti:
    # _hide_table_borders_keep_layout(stmt_tbl, mode="white")

    doc.add_paragraph("")
    tanggal = _format_tanggal(exec_dt)
    p = doc.add_paragraph(f"Banda Aceh,     {tanggal}\nYang menyatakan,")
    _style(p, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("")
    p_m = doc.add_paragraph("Materai 10000")
    _style(p_m, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for run in p_m.runs:
        if "Materai" in run.text:
            run.font.color.rgb = RGBColor(160, 160, 160)

    doc.add_paragraph("")
    p = doc.add_paragraph(f"{person.nama}\nNIP. {person.nip}")
    _style(p, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ===== HALAMAN 2 =====
    doc.add_page_break()

    p = doc.add_paragraph(
        "Lampiran Daftar Biaya Submit Artikel/Insentif Publikasi/Opini Media Massa/"
        f"Hak Kekayaan Intelektual yang didanai atas nama {person.nama} sebagai berikut:"
    )
    _style(p, justify=True)

    # Lampiran (tetap seperti awal, hanya tambah autofit False + set column width)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.autofit = False

    widths = [Cm(2.7), Cm(8.4), Cm(2.5), Cm(3.5)]
    headers = ["No. Proposal", "Judul Insentif", "Skema", "Jumlah Dana (Rp)"]

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        c.width = widths[i]
        _style(c.paragraphs[0], size=10, bold=True)

    for row in lampiran_rows:
        r = tbl.add_row().cells
        values = [
            str(row.get("no_prop", "")).strip(),
            str(row.get("judul", "")).strip(),
            str(row.get("skema", "")).strip(),
            str(row.get("dana", "")).strip(),
        ]
        for i, val in enumerate(values):
            r[i].text = val
            r[i].width = widths[i]
            _style(r[i].paragraphs[0], size=10)

    doc.add_paragraph("")
    doc.add_paragraph("")

    ttd = doc.add_table(rows=1, cols=3)
    ttd.autofit = False
    ttd.style = "Table Grid"
    ttd.columns[0].width = Cm(16 - 1.5 - 2.0)
    ttd.columns[1].width = Cm(1.5)
    ttd.columns[2].width = Cm(2.0)

    ttd.rows[0].cells[0].text = ""
    _remove_cell_borders(ttd.rows[0].cells[0])
    ttd.rows[0].cells[1].text = "Tanda\nTangan"
    _style(ttd.rows[0].cells[1].paragraphs[0], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    ttd.rows[0].cells[2].text = ""

    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


# =========================================================
# ITER PEOPLE FROM DF
# =========================================================
def iter_people_from_df(df: pd.DataFrame) -> Tuple[Person, List[Dict[str, str]]]:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    max_n = detect_max_n(df)

    bank_col = (
        "Nama Bank"
        if "Nama Bank" in df.columns
        else ("nama_bank" if "nama_bank" in df.columns else None)
    )
    email_col = (
        "Email"
        if "Email" in df.columns
        else ("email" if "email" in df.columns else None)
    )

    for _, row in df.iterrows():
        nip = str(row.get("NIP", "")).strip()
        nama = str(row.get("Nama", "")).strip()
        fakultas = str(row.get("Fakultas", "")).strip()
        rekening = str(row.get("Norek", "")).strip()

        if not nip or not nama:
            continue

        bank = str(row.get(bank_col, "")).strip() if bank_col else "-"
        email = normalize_email(row.get(email_col, None)) if email_col else None

        lampiran_rows: list[dict] = []
        for i in range(1, max_n + 1):
            no_prop = row.get(f"NoProp{i}", None)
            if pd.isna(no_prop) or str(no_prop).strip() == "":
                continue

            judul = row.get(f"Judul{i}", None)
            skema = row.get(f"Skema{i}", None)
            dana = row.get(f"Jumlah_dana{i}", None)

            lampiran_rows.append(
                {
                    "no_prop": str(no_prop).strip(),
                    "judul": "" if pd.isna(judul) else str(judul).strip(),
                    "skema": "" if pd.isna(skema) else str(skema).strip(),
                    "dana": _format_rupiah(dana),
                }
            )

        if not lampiran_rows:
            continue

        person = Person(
            nama=nama,
            nip=nip,
            fakultas=fakultas,
            rekening=rekening,
            bank=bank or "-",
            email=email,
        )
        yield person, lampiran_rows


# =========================================================
# GENERATE PDF PER PERSON
# =========================================================
def generate_pdf_bytes_for_person(
    person: Person,
    lampiran_rows: List[Dict[str, str]],
    *,
    exec_dt: Optional[datetime] = None,
    soffice_path: str | None = None,
) -> bytes:
    exec_dt = exec_dt or datetime.now()
    doc = _build_docx_from_scratch(person, lampiran_rows, exec_dt)

    docx_buf = io.BytesIO()
    doc.save(docx_buf)

    return convert_docx_bytes_to_pdf_bytes(
        docx_buf.getvalue(),
        soffice_path=soffice_path,
        timeout_sec=120,
    )


def filename_for_person(person: Person) -> str:
    return _safe_filename(person.nama, person.nip, "pdf")


def generate_zip_pdf(pdf_items: List[Tuple[str, bytes]]) -> bytes:
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as z:
        for filename, pdf_bytes in pdf_items:
            z.writestr(filename, pdf_bytes)
    zip_buf.seek(0)
    return zip_buf.read()
